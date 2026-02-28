#!/usr/bin/env python3
"""
Parse DOCX questions and generate SQL for organic chemistry topics.

Usage:
  python parse_tema1_docx_to_sql.py /path/to/file.docx
  python parse_tema1_docx_to_sql.py /path/to/file.docx --topic "Гетероциклічні сполуки 28,6" --images-subdir heterotsyklichni --header-lines 2

Outputs:
  - public/test-images/organic/{subdir}/img{N}.png (extracted images)
  - supabase/{output}.sql (INSERT statements)
"""

import argparse
import os
import re
import sys
import zipfile
from pathlib import Path

from docx import Document

COURSE_TITLE = "Основи знань про органічні сполуки"

# Defaults for Тема 1
DEFAULT_TOPIC = "ВСТУП. Основи знань про органічні сполуки"
DEFAULT_SUBDIR = "vstup"
DEFAULT_HEADER_LINES = 8


def has_drawing(para) -> bool:
    """Check if paragraph contains a drawing (image)."""
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    drawings = para._element.findall(".//w:drawing", ns)
    picts = para._element.findall(".//w:pict", ns)
    return len(drawings) > 0 or len(picts) > 0


def get_image_rids(para) -> list:
    """Extract rId values for images in paragraph."""
    rids = []
    for elem in para._element.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        if tag == "blip":
            rid = elem.get("{" + ns + "}embed")
            if rid:
                rids.append(rid)
        if tag == "imagedata":
            rid = elem.get("{" + ns + "}id")
            if rid:
                rids.append(rid)
    return rids


def normalize(text: str) -> str:
    """Replace non-breaking spaces, fix known typos, and normalize for SQL."""
    if not text:
        return ""
    text = text.replace("\xa0", " ").strip()
    # Fix P264: "полука" -> "Сполука" (missing first char in source)
    if text.startswith("полука відноситься до похідних"):
        text = "Сполука" + text[6:]  # skip "полука" (6 chars)
    return text


def sql_escape(text: str) -> str:
    """Escape single quotes for SQL."""
    return text.replace("'", "''")


def looks_like_question(text: str) -> bool:
    """Heuristic: question text is usually longer and ends with : or ?"""
    t = (text or "").strip()
    if len(t) < 15:
        return False
    return t.endswith(":") or t.endswith("?") or t.endswith(".")


def parse_questions(doc, rid_to_path: dict, header_lines: int = 8) -> list:
    """
    Parse document into list of {question_text, image_urls, options}.
    Images are associated with the question they follow in document flow.
    """
    questions = []
    paras = doc.paragraphs
    pos = header_lines
    pending_question_prefix = None

    def is_image_para(idx: int) -> bool:
        if idx >= len(paras):
            return False
        return has_drawing(paras[idx])

    while pos < len(paras):
        p = paras[pos]
        text = normalize(p.text)

        # Check for merged paragraph: "option@next_question_text" (e.g. P473)
        if "@" in text and not text.strip().endswith("@"):
            parts = text.split("@", 1)
            opt_part = (parts[0] + "@").strip()
            next_q = (parts[1] or "").strip()
            if opt_part and next_q:
                pending_question_prefix = next_q
                text = opt_part

        # Image paragraph - add to current (last) question
        # If we have pending_question_prefix, the image belongs to the NEXT question - start it first
        if is_image_para(pos):
            if pending_question_prefix and questions and len(questions[-1]["options"]) == 5:
                questions.append({
                    "question_text": pending_question_prefix,
                    "image_urls": [],
                    "options": [],
                })
                pending_question_prefix = None
            rids = get_image_rids(p)
            paths = [rid_to_path[r] for r in rids if r in rid_to_path]
            if questions and paths:
                questions[-1]["image_urls"].extend(paths)
            pos += 1
            continue

        # pending_question_prefix is used when we see an image (create new question) - never prepend to current text

        if looks_like_question(text) and (
            not questions or len(questions[-1]["options"]) == 5
        ):
            questions.append({
                "question_text": text,
                "image_urls": [],
                "options": [],
            })
            pos += 1
            continue

        # Option
        if questions and len(questions[-1]["options"]) < 5 and text:
            is_correct = "@" in text
            opt_text = text.replace("@", "").strip()
            questions[-1]["options"].append({
                "text": opt_text,
                "is_correct": is_correct,
            })
            pos += 1
            continue

        # Stray or new question after full block
        if questions and len(questions[-1]["options"]) == 5 and text:
            if looks_like_question(text):
                questions.append({
                    "question_text": text,
                    "image_urls": [],
                    "options": [],
                })
            pos += 1
        else:
            pos += 1

    return questions


def extract_images_in_order(
    docx_path: str, doc, output_dir: str, public_prefix: str
) -> dict:
    """
    Extract images in document order and build rid_to_path.
    Must be called before parse_questions so we have paths for image association.
    """
    os.makedirs(output_dir, exist_ok=True)
    rid_to_path = {}
    img_counter = [0]

    with zipfile.ZipFile(docx_path, "r") as zf:
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        rel_pattern = re.compile(
            r'Relationship Id="(rId\d+)"[^>]*Target="([^"]+)"'
        )
        rels = {m.group(1): m.group(2) for m in rel_pattern.finditer(rels_xml)}

        def get_path_for_rid(rid: str) -> str:
            if rid in rid_to_path:
                return rid_to_path[rid]
            target = rels.get(rid, "")
            if not target or "image" not in target.lower():
                return ""
            try:
                data = zf.read("word/" + target)
            except KeyError:
                return ""
            ext = os.path.splitext(target)[1] or ".png"
            if ext.lower() not in (".png", ".jpg", ".jpeg", ".gif"):
                ext = ".png"
            img_counter[0] += 1
            out_name = f"img{img_counter[0]}{ext}"
            out_path = os.path.join(output_dir, out_name)
            with open(out_path, "wb") as f:
                f.write(data)
            public_path = f"{public_prefix}/{out_name}"
            rid_to_path[rid] = public_path
            return public_path

        for para in doc.paragraphs:
            if has_drawing(para):
                for rid in get_image_rids(para):
                    get_path_for_rid(rid)

    return rid_to_path


def generate_sql(
    questions: list, sql_path: str, topic_title: str, course_title: str = COURSE_TITLE
) -> None:
    """Generate SQL file."""
    topic_esc = sql_escape(topic_title)
    lines = [
        "BEGIN;",
        "",
        f"-- Тема: {topic_title}",
        "",
        "INSERT INTO topics (course_id, title, description, order_index, is_active)",
        f"SELECT c.id, '{topic_esc}', '{topic_esc}', 0, true",
        f"FROM courses c",
        f"WHERE c.title = '{course_title}'",
        f"  AND NOT EXISTS (SELECT 1 FROM topics t WHERE t.course_id = c.id AND t.title = '{topic_esc}');",
        "",
        "DELETE FROM question_options WHERE question_id IN (",
        "  SELECT q.id FROM questions q",
        "  JOIN topics t ON t.id = q.topic_id",
        "  JOIN courses c ON c.id = t.course_id",
        f"  WHERE c.title = '{course_title}'",
        f"    AND t.title = '{topic_esc}'",
        ");",
        "",
        "DELETE FROM questions WHERE topic_id = (",
        "  SELECT t.id FROM topics t",
        "  JOIN courses c ON c.id = t.course_id",
        f"  WHERE c.title = '{course_title}'",
        f"    AND t.title = '{topic_esc}'",
        ");",
        "",
    ]

    for idx, q in enumerate(questions):
        order_idx = idx + 1
        qtext = sql_escape(q["question_text"])
        img_urls = q.get("image_urls", [])
        image_url = ",".join(img_urls) if img_urls else "NULL"
        if image_url and image_url != "NULL":
            image_url = "'" + image_url + "'"

        opts = q.get("options", [])

        lines.append(f"-- Question {order_idx}")
        lines.append("WITH inserted_question AS (")
        lines.append(
            "  INSERT INTO questions (topic_id, question_text, explanation, difficulty, order_index, is_active, image_url)"
        )
        lines.append(
            f"  SELECT t.id, '{qtext}', NULL, 'medium', {order_idx}, true, {image_url}"
        )
        lines.append("  FROM topics t")
        lines.append("  JOIN courses c ON c.id = t.course_id")
        lines.append(f"  WHERE c.title = '{course_title}'")
        lines.append(f"    AND t.title = '{topic_esc}'")
        lines.append("  RETURNING id")
        lines.append(")")
        opt_rows = []
        for oi, opt in enumerate(opts):
            otext = sql_escape(opt["text"])
            corr = "true" if opt.get("is_correct") else "false"
            opt_rows.append(
                f"  ((SELECT id FROM inserted_question), '{otext}', {corr}, {oi})"
            )
        lines.append("INSERT INTO question_options (question_id, option_text, is_correct, order_index) VALUES")
        lines.append(",\n".join(opt_rows) + ";")
        lines.append("")
        lines.append("")

    lines.append("COMMIT;")

    with open(sql_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def main():
    parser = argparse.ArgumentParser(
        description="Parse DOCX questions and generate SQL for organic chemistry topics."
    )
    parser.add_argument("docx_path", help="Path to .docx file")
    parser.add_argument(
        "--topic",
        default=DEFAULT_TOPIC,
        help=f"Topic title (default: {DEFAULT_TOPIC})",
    )
    parser.add_argument(
        "--images-subdir",
        default=DEFAULT_SUBDIR,
        help=f"Subdir under public/test-images/organic/ (default: {DEFAULT_SUBDIR})",
    )
    parser.add_argument(
        "--header-lines",
        type=int,
        default=DEFAULT_HEADER_LINES,
        help=f"Paragraphs to skip at start (default: {DEFAULT_HEADER_LINES})",
    )
    parser.add_argument(
        "--sql-output",
        default=None,
        help="Output SQL filename (default: derived from subdir)",
    )
    parser.add_argument(
        "--course",
        default=COURSE_TITLE,
        help=f"Course title in DB (default: {COURSE_TITLE!r})",
    )
    parser.add_argument(
        "--images-base",
        default="organic",
        help="Base dir under public/test-images/ for images (default: organic)",
    )
    args = parser.parse_args()

    docx_path = args.docx_path
    if not os.path.isfile(docx_path):
        print(f"File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    script_dir = Path(__file__).parent.resolve()
    images_dir = script_dir / "public" / "test-images" / args.images_base / args.images_subdir
    public_prefix = f"/test-images/{args.images_base}/{args.images_subdir}"
    sql_name = args.sql_output or f"{args.images_subdir}-questions.sql"
    sql_path = script_dir / "supabase" / sql_name

    print("Loading document...", file=sys.stderr)
    doc = Document(docx_path)

    print("Extracting images...", file=sys.stderr)
    rid_to_path = extract_images_in_order(
        docx_path, doc, str(images_dir), public_prefix
    )
    print(f"  Extracted {len(rid_to_path)} images to {images_dir}", file=sys.stderr)

    print("Parsing questions...", file=sys.stderr)
    questions = parse_questions(doc, rid_to_path, args.header_lines)
    print(f"  Found {len(questions)} questions", file=sys.stderr)

    print("Generating SQL...", file=sys.stderr)
    generate_sql(questions, str(sql_path), args.topic, args.course)
    print(f"  Written to {sql_path}", file=sys.stderr)
    print("Done.", file=sys.stderr)


if __name__ == "__main__":
    main()
